from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import re
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCX_SOURCE_FILES = [
    "docs/CROSSFRAME_V6.md",
    "docs/V5_TO_V6_CHANGES.md",
    "docs/V6_TOOL_PROTOTYPE.md",
]

TOOL_BOUNDARY = "本工具只帮助整理结构判断，不替代现实调查、专业判断、受影响位置反馈、外部复核和最终责任承担。分数只能触发降级、复核或补证，不能单独授权处置。"

DOCX_PRESET_NAME = "compact_reference_guide"
BASE_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
FIXED_DOCX_TIMESTAMP = (2000, 1, 1, 0, 0, 0)

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
INK = RGBColor(0x20, 0x20, 0x20)


def set_run_font(
    run,
    *,
    font: str = BASE_FONT,
    east_asia: str = EAST_ASIA_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, font: str, east_asia: str, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    style.font.name = font
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    set_style_font(normal, font=BASE_FONT, east_asia=EAST_ASIA_FONT, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = document.styles[style_name]
        set_style_font(style, font=BASE_FONT, east_asia=EAST_ASIA_FONT, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = document.styles[style_name]
        set_style_font(style, font=BASE_FONT, east_asia=EAST_ASIA_FONT, size=11, color=INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("跨尺度结构诊断框架 v6.0")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("CrossFrame v6.0 · 半量化审计层")
    set_run_font(run, size=9, color=MUTED)


def add_inline_runs(paragraph, text: str, *, size: float | None = None, color: RGBColor | None = None) -> None:
    pos = 0
    bold = False
    while pos < len(text):
        if text.startswith("**", pos):
            end = text.find("**", pos + 2)
            if end != -1:
                content = text[pos + 2:end]
                run = paragraph.add_run(content)
                set_run_font(run, size=size, color=color, bold=True)
                pos = end + 2
                continue
        if text[pos] == "`":
            end = text.find("`", pos + 1)
            if end != -1:
                content = text[pos + 1:end]
                run = paragraph.add_run(content)
                set_run_font(run, font=MONO_FONT, east_asia=EAST_ASIA_FONT, size=size or 9.5, color=color)
                pos = end + 1
                continue
        run = paragraph.add_run(text[pos])
        set_run_font(run, size=size, color=color, bold=bold)
        pos += 1


def add_paragraph(document: Document, text: str, *, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    add_inline_runs(paragraph, text)


def add_code_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.right_indent = Inches(0.1)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(line if line else " ")
        set_run_font(run, font=MONO_FONT, east_asia=EAST_ASIA_FONT, size=9.2, color=RGBColor(0x33, 0x33, 0x33))


def source_digest(repo: Path) -> str:
    digest = hashlib.sha256()
    for rel in DOCX_SOURCE_FILES:
        digest.update(rel.encode("utf-8"))
        digest.update(b"\0")
        digest.update((repo / rel).read_bytes())
        digest.update(b"\0")
    return digest.hexdigest()


def normalize_docx_zip(path: Path) -> None:
    original = path.read_bytes()
    tmp_path = path.with_suffix(path.suffix + ".tmp")
    with ZipFile(path, "r") as src, ZipFile(tmp_path, "w", compression=ZIP_DEFLATED) as dst:
        for name in sorted(src.namelist()):
            src_info = src.getinfo(name)
            info = ZipInfo(name, FIXED_DOCX_TIMESTAMP)
            info.compress_type = ZIP_DEFLATED
            info.external_attr = src_info.external_attr
            data = src.read(name)
            dst.writestr(info, data)
    if tmp_path.read_bytes() == original:
        tmp_path.unlink()
        return
    try:
        tmp_path.replace(path)
    except PermissionError:
        path.unlink()
        tmp_path.replace(path)


def is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
    except ValueError:
        return False
    return True


def validate_generated_output_target(repo: Path, target: Path) -> None:
    resolved = target.resolve()
    repo_resolved = repo.resolve()
    if not is_relative_to(resolved, repo_resolved):
        return
    outputs_root = (repo_resolved / "outputs").resolve()
    if not is_relative_to(resolved, outputs_root):
        raise SystemExit(
            f"refusing to write generated DOCX inside tracked repo tree outside outputs/: {resolved}"
        )


def add_title_page(document: Document, source_hash: str, release_label: str | None) -> None:
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("跨尺度结构诊断框架 v6.0")
    set_run_font(run, size=25, color=RGBColor(0x00, 0x00, 0x00), bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("半量化结构判断系统发布稿")
    set_run_font(run, size=14, color=MUTED)

    metadata = [
        ("生成标签", release_label or "source-digest"),
        ("源稿", ", ".join(DOCX_SOURCE_FILES)),
        ("源稿摘要", source_hash[:16]),
        ("文档预设", DOCX_PRESET_NAME),
        ("定位", "v5.0 结构诊断协议之上的半量化审计层"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=MUTED)

    document.add_paragraph()
    add_callout(document, TOOL_BOUNDARY)
    document.add_page_break()


def add_static_toc(document: Document, headings: list[str]) -> None:
    paragraph = document.add_paragraph("目录", style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    for heading in headings:
        item = document.add_paragraph(style="List Bullet")
        add_inline_runs(item, heading)
    document.add_page_break()


def add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    apply_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=RGBColor(0x1F, 0x3A, 0x5F), bold=True)
    document.add_paragraph()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side, value in CELL_MARGIN_DXA.items():
        node = tbl_cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(table)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def weighted_widths(rows: list[list[str]]) -> list[int]:
    col_count = len(rows[0])
    if col_count == 1:
        return [CONTENT_WIDTH_DXA]
    if col_count == 2:
        return [2700, CONTENT_WIDTH_DXA - 2700]
    if col_count == 3:
        return [2000, 3680, 3680]
    if col_count == 4:
        return [1600, 2600, 2600, 2560]
    if col_count == 5:
        return [1400, 1750, 2550, 1700, 1960]

    base = CONTENT_WIDTH_DXA // col_count
    widths = [base for _ in range(col_count)]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = weighted_widths(rows)
    apply_table_geometry(table, widths)
    set_repeat_header(table.rows[0])

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
            run_size = 9.2 if col_count >= 5 else 9.8
            add_inline_runs(paragraph, value, size=run_size)
            for run in paragraph.runs:
                run.bold = row_idx == 0
    document.add_paragraph()


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [part.strip() for part in stripped.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def collect_headings(markdown: str, *, shift: int, skip_first_h1: bool) -> list[str]:
    headings: list[str] = []
    skipped = False
    for line in markdown.splitlines():
        match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if not match:
            continue
        level = len(match.group(1))
        if level == 1 and skip_first_h1 and not skipped:
            skipped = True
            continue
        style_level = max(1, min(3, level + shift))
        if style_level == 1:
            headings.append(match.group(2).strip())
    return headings


def render_markdown(document: Document, markdown: str, *, shift: int, skip_first_h1: bool) -> None:
    lines = markdown.splitlines()
    i = 0
    skipped_h1 = False
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if i + 1 < len(lines) and "|" in line and is_table_separator(lines[i + 1]):
            table_rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and "|" in lines[i].strip():
                table_rows.append(split_table_row(lines[i]))
                i += 1
            add_markdown_table(document, table_rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1 and skip_first_h1 and not skipped_h1:
                skipped_h1 = True
                i += 1
                continue
            style_level = max(1, min(3, level + shift))
            document.add_paragraph(text, style=f"Heading {style_level}")
            i += 1
            continue

        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet.group(1).strip())
            i += 1
            continue

        number = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if number:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, number.group(1).strip())
            i += 1
            continue

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        add_paragraph(document, line.strip())
        i += 1

    if code_lines:
        add_code_block(document, code_lines)


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def build_docx(repo: Path, output_path: Path, release_label: str | None) -> None:
    missing = [rel for rel in DOCX_SOURCE_FILES if not (repo / rel).exists()]
    if missing:
        raise SystemExit("missing DOCX source file(s): " + ", ".join(missing))
    validate_generated_output_target(repo, output_path)
    source_hash = source_digest(repo)

    document = Document()
    configure_styles(document)
    set_header_footer(document)
    document.core_properties.title = "跨尺度结构诊断框架 v6.0"
    document.core_properties.subject = "半量化结构判断系统发布稿"
    document.core_properties.author = "CrossFrame Skill Suite"
    document.core_properties.keywords = "CrossFrame, v6.0, 半量化, 结构诊断"
    document.core_properties.created = dt.datetime(*FIXED_DOCX_TIMESTAMP)
    document.core_properties.modified = dt.datetime(*FIXED_DOCX_TIMESTAMP)
    document.core_properties.revision = 1

    add_title_page(document, source_hash, release_label)

    main_text = read(repo / DOCX_SOURCE_FILES[0])
    toc_headings = collect_headings(main_text, shift=-1, skip_first_h1=True)
    toc_headings.extend(["附录 A：v5.0 到 v6.0 变更说明", "附录 B：v6.0 工具原型说明"])
    add_static_toc(document, toc_headings)

    render_markdown(document, main_text, shift=-1, skip_first_h1=True)

    appendix_a = read(repo / DOCX_SOURCE_FILES[1])
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("附录 A：v5.0 到 v6.0 变更说明", style="Heading 1")
    render_markdown(document, appendix_a, shift=0, skip_first_h1=True)

    appendix_b = read(repo / DOCX_SOURCE_FILES[2])
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("附录 B：v6.0 工具原型说明", style="Heading 1")
    render_markdown(document, appendix_b, shift=0, skip_first_h1=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    normalize_docx_zip(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the CrossFrame v6.0 publishable DOCX from tracked Markdown sources.")
    parser.add_argument("--repo", default=".", help="Repository root.")
    parser.add_argument("--output-dir", default="outputs/doc", help="Directory for generated DOCX output.")
    parser.add_argument("--filename", default="跨尺度结构诊断框架v6.0.docx", help="Output DOCX filename.")
    parser.add_argument("--date", default=None, help="Optional release/date label. Defaults to source-digest metadata.")
    parser.add_argument("--copy-to", default=None, help="Optional additional DOCX path to write.")
    args = parser.parse_args()

    repo = Path(args.repo).resolve()
    release_label = args.date
    output_path = repo / args.output_dir / args.filename
    build_docx(repo, output_path, release_label)
    print(f"created: {output_path}")
    if args.copy_to:
        copy_path = Path(args.copy_to).resolve()
        build_docx(repo, copy_path, release_label)
        print(f"created: {copy_path}")
    print(f"sources: {len(DOCX_SOURCE_FILES)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
