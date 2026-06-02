from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document


BUNDLE_DEFS = {
    "framework-use-discipline-pack": {
        "name": "框架使用纪律包",
        "summary": "约束 CrossFrame 如何使用：防概念武器化、防教条化、防万能化、防 AI 文本替代证据。",
        "must": "必须联读使用准则、输出禁忌、本体保护、AI 输出分离和适用性分级。",
    },
    "judgment-responsibility-pack": {
        "name": "判断责任包",
        "summary": "约束开放断言、强判断、命题验证和高责任处置，防止把可撤回判断当终局裁决。",
        "must": "必须联读证据成本、判断档位、命题验证、申诉/反证入口、弱信号保护和反俘获。",
    },
    "diagnosis-mainline-pack": {
        "name": "诊断主线包",
        "summary": "保持对象、事实、尺度、机制候选、五闸、工具箱和诊断维度的连续主线。",
        "must": "必须联读有限干预、观测登记、诊断档位、主流程、五闸十三步、诊断工具箱和维度。",
    },
    "intimate-love-care-pack": {
        "name": "亲密关系/爱/照护包",
        "summary": "处理关系、照护、爱、解释劳动和疗愈时，先保护痛苦、安全和边界。",
        "must": "必须联读亲密关系轻量入口、爱诊断红线、开放行动、修复副产品、责任链和疗愈/转移。",
    },
    "public-power-governance-pack": {
        "name": "公共制度与权力包",
        "summary": "处理平台、制度、公共承诺、程序有效性、权力封闭和低权力主体保护。",
        "must": "必须联读反俘获、程序有效性、弱信号安全、退出转移、公共承诺、偿付和多中心治理。",
    },
    "long-evolution-deep-pack": {
        "name": "长期演化深水区包",
        "summary": "处理根假设、生命周期、递进、势场、自主解离、治理连续性和文明尺度压力测试。",
        "must": "必须联读根假设/核心推论、生命周期、递进、势场、解离、调节预警偿付、多中心治理和超大规模压力测试。",
    },
    "expression-article-pack": {
        "name": "表达与文章输出包",
        "summary": "把后台概念翻译成人话、管理/制度/技术语境或文章输出，避免术语墙。",
        "must": "必须联读对外表达翻译、AI 通俗模板、用户语言闸和文章底稿/正文规则。",
    },
}

BUNDLE_ORDER = list(BUNDLE_DEFS)


def has_any(text: str, keywords: list[str]) -> bool:
    return any(keyword in text for keyword in keywords)


def bundles_for(title: str) -> list[str]:
    if "目录" == title:
        return ["source-structure-only"]

    bundles: set[str] = set()
    if has_any(
        title,
        [
            "极简导读",
            "模型",
            "什么时候",
            "入口概念",
            "轻量诊断流程",
            "常见误用",
            "框架定位",
            "解释对象",
            "跨尺度迁移闸",
            "使用准则",
            "概念武器化",
            "教条化",
            "过度自缚",
            "十二条",
            "信号成本",
            "弱信号",
            "观测参与",
            "反身性准则",
            "框架本体保护",
            "反领域",
            "反模型",
            "适用性",
            "输出禁忌",
            "模块化入口",
            "操作层通则",
            "诊断与干预分离",
            "责任归属",
            "趋势判断限制",
            "高风险概念边界",
            "盲区声明",
            "概念层级",
            "概念保真",
            "前台-后台",
            "修改资格",
            "概念融合",
            "边界元假设",
            "根假设使用纪律",
        ],
    ):
        bundles.add("framework-use-discipline-pack")
    if has_any(
        title,
        [
            "强判断",
            "开放断言",
            "命题验证",
            "前瞻登记",
            "判断责任",
            "反向条件",
            "修复窗口",
            "证据要求",
            "申诉入口",
            "开放断言许可",
            "判断勇气",
            "证伪",
            "反叙事",
            "反例压力",
            "高责任",
            "权力封闭",
            "低条件",
            "试探行动",
            "诊断可行性",
            "反俘获",
            "程序有效性",
            "AI 合规",
            "渐进式诊断契约",
            "高反身性博弈",
        ],
    ):
        bundles.add("judgment-responsibility-pack")
    if has_any(
        title,
        [
            "诊断",
            "预判",
            "有限干预",
            "观测影响",
            "诊断档位",
            "诊断主流程",
            "圈层界定",
            "核心分析对象",
            "主导约束",
            "演化阶段",
            "锚点",
            "先行者",
            "推力链",
            "主体层",
            "失稳因素",
            "长期隐患",
            "分支预判",
            "完备性",
            "五闸",
            "工具箱",
            "中层承接器",
            "高频小回流",
            "抽样稽核",
            "摘要失真",
            "注意力分诊",
            "主动收束",
            "生命节点",
            "低破坏争议",
            "平庸化责任",
            "尺度升维",
            "跨圈层",
            "舆情",
            "边界类型",
            "局部排除区",
            "机制候选",
            "非线性闭合",
            "有效对象",
            "基础诊断维度",
            "扩展诊断维度",
            "锚点健康度",
            "动力余量",
            "递进健康度",
            "疗愈准备度",
        ],
    ):
        bundles.add("diagnosis-mainline-pack")
    if has_any(
        title,
        [
            "亲密",
            "爱",
            "照护",
            "家庭",
            "解释劳动",
            "疗愈",
            "抢救",
            "修复",
            "高阶稳态",
            "牺牲",
            "开放性承担",
            "不浪费爱",
            "爱诊断",
            "结构浪费",
            "生命节点承接",
        ],
    ):
        bundles.add("intimate-love-care-pack")
    if has_any(
        title,
        [
            "公共",
            "制度",
            "治理",
            "权力",
            "平台",
            "合规",
            "申诉",
            "承诺",
            "偿付",
            "多中心",
            "承接者",
            "代际",
            "生产",
            "分配",
            "再生产",
            "社会制度",
            "高权力密度",
            "低权力主体",
            "策略行动",
            "解释锚",
            "程序有效性",
            "AI 合规",
            "权力封闭",
            "调节",
            "预警",
        ],
    ):
        bundles.add("public-power-governance-pack")
    if has_any(
        title,
        [
            "理论框架",
            "核心概念",
            "锚点组",
            "动力组",
            "结构组",
            "过程组",
            "根假设",
            "核心推论",
            "全周期",
            "生命周期",
            "阶段",
            "递进",
            "双向势场",
            "势场",
            "自主解离",
            "沉积",
            "基本盘",
            "超大规模",
            "文明",
            "历史",
            "多层条件场",
            "高尺度环境",
            "结构负荷",
            "尺度升维下",
            "跨域互操作",
            "多层锚",
            "启动-转译",
            "支撑通道",
            "行动承接",
            "中层承接",
            "失稳机制",
            "过程组",
            "时间不可逆",
            "熵增",
            "嵌套耦合",
        ],
    ):
        bundles.add("long-evolution-deep-pack")
    if has_any(title, ["输出", "表达", "翻译", "模板", "格式", "AI 通俗", "说人话", "对外表达", "语境", "文章", "发布"]):
        bundles.add("expression-article-pack")

    if not bundles:
        bundles.add("diagnosis-mainline-pack")
    return [bundle for bundle in BUNDLE_ORDER if bundle in bundles]


def status_for(title: str, bundles: list[str]) -> str:
    if "source-structure-only" in bundles:
        return "源结构节点"
    if has_any(title, ["协议", "流程", "入口", "模板", "记录表", "命题验证表", "规则", "检查", "诊断主流程", "五闸", "强判断四格", "疗愈方案", "输出格式"]):
        return "已协议化/已联读约束"
    if has_any(title, ["概念", "锚点组", "动力组", "结构组", "过程组", "证据成本", "机制候选", "修复副产品", "责任链", "爱", "退出转移", "权力封闭", "反身性"]):
        return "已概念卡化/已联读约束"
    if has_any(title, ["根假设", "核心推论", "理论", "全周期", "递进", "势场", "自主解离", "多中心", "超大规模"]):
        return "已索引化/后台联读"
    return "已联读约束"


def carry_for(bundles: list[str]) -> str:
    if "source-structure-only" in bundles:
        return "references/v2-source-spine.md"
    mapping = {
        "framework-use-discipline-pack": "references/guardrails.md；references/framework-ontology-protection.md；protocols/framework-boundary-protocol.md",
        "judgment-responsibility-pack": "protocols/open-assertion-protocol.md；protocols/proposition-verification-protocol.md；protocols/anti-capture-protocol.md",
        "diagnosis-mainline-pack": "protocols/diagnosis-protocol.md；worksheets/five-gates-worksheet.md；references/diagnostic-dimensions.md；references/diagnostic-toolbox-index.md",
        "intimate-love-care-pack": "protocols/intimate-relationship-protocol.md；protocols/healing-transfer-protocol.md；references/concept-cards/love-open-action.md",
        "public-power-governance-pack": "protocols/public-institution-protocol.md；protocols/anti-capture-protocol.md；protocols/governance-continuity-protocol.md",
        "long-evolution-deep-pack": "references/theory-backend-index.md；protocols/lifecycle-diagnosis-protocol.md；protocols/progression-protocol.md；protocols/field-dissociation-protocol.md",
        "expression-article-pack": "protocols/expression-translation-protocol.md；templates/user-facing-language.md；../crossframe-essay/SKILL.md",
    }
    return "；".join(mapping[bundle] for bundle in bundles if bundle in mapping)


def digest_for(title: str, bundles: list[str]) -> str:
    names = [BUNDLE_DEFS[bundle]["name"] for bundle in bundles if bundle in BUNDLE_DEFS]
    if not names:
        return "源结构目录节点，用于保持 2.0 原文顺序，不承担独立判断。"
    return f"本节归入{'、'.join(names)}，主要用于保持“{title}”在原文中的连续约束；读取时要把局部概念放回相邻章节和联读包中理解。"


def boundary_for(bundles: list[str]) -> str:
    if "source-structure-only" in bundles:
        return "不能把目录节点当概念依据。"
    parts: list[str] = []
    if "framework-use-discipline-pack" in bundles:
        parts.append("不能把框架当万能审判或专业替代。")
    if "judgment-responsibility-pack" in bundles:
        parts.append("不能用开放断言替代强判断验证。")
    if "diagnosis-mainline-pack" in bundles:
        parts.append("不能跳过对象、证据、尺度、责任和观测闸。")
    if "intimate-love-care-pack" in bundles:
        parts.append("不能把爱或照护写成单方忍耐义务。")
    if "public-power-governance-pack" in bundles:
        parts.append("不能把程序外观或 AI 合规文本当强证据。")
    if "long-evolution-deep-pack" in bundles:
        parts.append("不能把阶段/势场/文明尺度写成宿命结论。")
    if "expression-article-pack" in bundles:
        parts.append("不能用术语墙替代普通人可读表达。")
    return "".join(parts)


def extract_headings(docx_path: Path) -> tuple[int, list[dict[str, object]]]:
    doc = Document(str(docx_path))
    headings: list[dict[str, object]] = []
    for idx, para in enumerate(doc.paragraphs):
        text = " ".join(para.text.split())
        style = para.style.name
        if text and style.startswith("Heading"):
            match = re.search(r"(\d+)$", style)
            level = int(match.group(1)) if match else 1
            headings.append({"para": idx, "level": level, "title": text})
    return sum(1 for para in doc.paragraphs if para.text.strip()), headings


def write_source_files(repo: Path, docx_path: Path) -> None:
    nonempty, headings = extract_headings(docx_path)
    for idx, heading in enumerate(headings):
        heading["id"] = f"V2-H{idx + 1:03d}"
        heading["prev"] = f"V2-H{idx:03d}" if idx else "-"
        heading["next"] = f"V2-H{idx + 2:03d}" if idx + 1 < len(headings) else "-"
        heading["bundles"] = bundles_for(str(heading["title"]))
        heading["status"] = status_for(str(heading["title"]), heading["bundles"])  # type: ignore[arg-type]
        heading["carry"] = carry_for(heading["bundles"])  # type: ignore[arg-type]
        heading["digest"] = digest_for(str(heading["title"]), heading["bundles"])  # type: ignore[arg-type]
        heading["boundary"] = boundary_for(heading["bundles"])  # type: ignore[arg-type]

    source = repo / "skills/crossframe/references/v2-source-spine.md"
    digest = repo / "skills/crossframe/references/v2-section-digest-index.md"

    source_lines = [
        "# v2.0 源结构脊柱\n",
        "\n",
        "本文件从 `跨尺度结构诊断框架v2.0.docx` 的 Word 标题层级生成，用于保存 2.0 原文的章节顺序、相邻关系和 CrossFrame 承接位置。它不是原文复制件；不得把本文件当作完整理论正文。\n",
        "\n",
        f"- 来源：`{docx_path.name}`\n",
        f"- 非空段落：{nonempty}\n",
        f"- Word 标题节点：{len(headings)}\n",
        "- 稳定 ID：`V2-H001` 起按原文标题出现顺序编号。\n",
        "- 使用方法：先按 `read-routing-map.md` 确定任务路由，再用本文件确认相邻章节和连续联读包。\n",
        "\n",
        "## 连续联读包图例\n",
        "\n",
        "| 联读包 ID | 中文名 | 作用 | 必须避免的读法 |\n",
        "| --- | --- | --- | --- |\n",
    ]
    for bundle_id, bundle in BUNDLE_DEFS.items():
        source_lines.append(
            f"| `{bundle_id}` | {bundle['name']} | {bundle['summary']} | 不得只读单个概念卡；{bundle['must']} |\n"
        )
    source_lines.extend(
        [
            "| `source-structure-only` | 源结构节点 | 目录或结构占位，只保持原文顺序 | 不得作为诊断依据 |\n",
            "\n",
            "## 源章节顺序表\n",
            "\n",
            "| ID | 段落 | 层级 | 标题 | 连续联读包 | 承接状态 | CrossFrame 承接位置 | 前一节 | 后一节 |\n",
            "| --- | ---: | ---: | --- | --- | --- | --- | --- | --- |\n",
        ]
    )
    for heading in headings:
        bundles = ", ".join(f"`{bundle}`" for bundle in heading["bundles"])  # type: ignore[index]
        source_lines.append(
            f"| `{heading['id']}` | P{heading['para']:04d} | H{heading['level']} | {heading['title']} | {bundles} | {heading['status']} | {heading['carry']} | `{heading['prev']}` | `{heading['next']}` |\n"
        )
    source.write_text("".join(source_lines), encoding="utf-8")

    digest_lines = [
        "# v2.0 逐节保真摘要索引\n",
        "\n",
        "本文件配合 `v2-source-spine.md` 使用。它不复刻 2.0 原文，而是给每个源章节一个保真用途摘要、不可误读边界和相邻联读提醒，用于防止拆成 skill 后读少、断章或误读。\n",
        "\n",
        "## 使用规则\n",
        "\n",
        "- 本文件只提供逐节摘要和保真提醒，不替代原文。\n",
        "- 当某节进入判断依据时，必须同时读取其连续联读包。\n",
        "- 若只读取单概念卡无法保留相邻约束，判断必须降档或补读。\n",
        "\n",
        "## 逐节索引\n",
        "\n",
        "| ID | 标题 | 保真用途摘要 | 不可误读边界 | 相邻联读提醒 |\n",
        "| --- | --- | --- | --- | --- |\n",
    ]
    for heading in headings:
        bundles = ", ".join(f"`{bundle}`" for bundle in heading["bundles"])  # type: ignore[index]
        digest_lines.append(
            f"| `{heading['id']}` | {heading['title']} | {heading['digest']} | {heading['boundary']} | 前接 `{heading['prev']}`，后接 `{heading['next']}`；同包：{bundles} |\n"
        )
    digest.write_text("".join(digest_lines), encoding="utf-8")


def update_read_routing(repo: Path) -> None:
    path = repo / "skills/crossframe/references/read-routing-map.md"
    text = path.read_text(encoding="utf-8")
    high_risk = text[text.index("## 高风险概念触发") :]
    intro = (
        "# 读取路由图\n\n"
        "本文件决定一次 CrossFrame 调用应该读取哪些材料。默认先读最少必要内容；当概念承担判断作用时，再加载完整概念卡和保真材料。\n\n"
        "本文件现在同时承担“连续联读”路由：当某个概念属于 2.0 原文连续板块时，不得只读单个 protocol 或 concept card，必须按 `continuity-bundles.md` 读取对应联读包，并在输出前使用 `worksheets/source-continuity-check.md` 检查是否读少。\n\n"
        "## 基础路由\n\n"
        "| 联读包 ID | 中文名 | 何时强制读取 |\n"
        "| --- | --- | --- |\n"
    )
    for bundle_id, bundle in BUNDLE_DEFS.items():
        intro += f"| `{bundle_id}` | {bundle['name']} | {bundle['summary']} |\n"
    intro += "\n"
    rows = [
        ("快速诊断", "`protocols/diagnosis-protocol.md`、`worksheets/intake-worksheet.md`、`worksheets/five-gates-worksheet.md`、`templates/quick-diagnosis-output.md`", "涉及高风险概念时读对应概念卡", "`diagnosis-mainline-pack`"),
        ("完整诊断 / 审计 / 深度分析", "`protocols/diagnosis-protocol.md`、全部核心 worksheets、`templates/full-diagnosis-output.md`、`references/v2-term-fidelity.md`", "`references/diagnostic-dimensions.md`、`references/diagnostic-toolbox-index.md`、`references/v2-source-spine.md`、`references/v2-section-digest-index.md`", "`diagnosis-mainline-pack`、`framework-use-discipline-pack`"),
        ("推演 / 后续走向 / 分支终点", "`protocols/inference-protocol.md`、`templates/inference-output.md`、`references/concept-cards/mechanism-candidates.md`", "尺度、反身性、权力封闭相关概念卡", "`diagnosis-mainline-pack`"),
        ("开放断言", "`protocols/open-assertion-protocol.md`、`worksheets/open-assertion-record.md`、`templates/open-assertion-output.md`、`references/concept-cards/open-assertion.md`", "`references/concept-cards/evidence-cost.md`、`references/concept-cards/judgment-grades.md`", "`judgment-responsibility-pack`"),
        ("低条件行动", "`protocols/low-condition-action-protocol.md`、`references/concept-cards/low-condition-action.md`", "`references/concept-cards/evidence-cost.md`、`references/concept-cards/judgment-grades.md`", "`judgment-responsibility-pack`"),
        ("高责任 / 权力密度 / 处分 / 名誉 / 权利 / 公共记忆", "`protocols/anti-capture-protocol.md`、`worksheets/high-responsibility-check.md`、`references/concept-cards/power-closure.md`、`references/concept-cards/evidence-cost.md`、`references/concept-cards/judgment-grades.md`", "`references/concept-cards/exit-transfer.md`、`references/diagnostic-toolbox-index.md`、`references/v2-source-spine.md`", "`judgment-responsibility-pack`、`public-power-governance-pack`、`framework-use-discipline-pack`"),
        ("强判断 / 处置依据 / 资格名誉资源权利", "`protocols/proposition-verification-protocol.md`、`worksheets/proposition-verification.md`、`worksheets/prospective-registration.md`、`templates/strong-judgment-output.md`", "`protocols/anti-capture-protocol.md`、`references/concept-cards/judgment-grades.md`", "`judgment-responsibility-pack`、`framework-use-discipline-pack`"),
        ("高反身性 / 表演 / 反制 / 研究诊断规则", "`protocols/high-reflexivity-protocol.md`、`worksheets/reflexivity-state-transfer.md`、`templates/high-reflexivity-output.md`", "`references/concept-cards/reflexivity.md`、`references/concept-cards/evidence-cost.md`", "`judgment-responsibility-pack`、`framework-use-discipline-pack`"),
        ("亲密关系 / 家庭 / 照护 / 解释劳动 / 单方承接", "`protocols/intimate-relationship-protocol.md`、`worksheets/intimate-relationship-light-check.md`、`templates/intimate-relationship-output.md`", "`references/concept-cards/love-open-action.md`、`references/concept-cards/repair-byproduct.md`、`references/concept-cards/responsibility-chain.md`、`references/v2-section-digest-index.md`", "`intimate-love-care-pack`、`diagnosis-mainline-pack`"),
        ("疗愈 / 修复路线 / 退出转移 / 长期重建", "`protocols/healing-transfer-protocol.md`、`worksheets/healing-transfer-map.md`、`templates/healing-transfer-output.md`", "`references/concept-cards/exit-transfer.md`、`references/concept-cards/repair-byproduct.md`", "`intimate-love-care-pack`、`diagnosis-mainline-pack`"),
        ("公共制度 / 平台治理 / 公共承诺 / 高权力公共议题", "`protocols/public-institution-protocol.md`、`worksheets/public-institution-check.md`、`templates/public-institution-output.md`", "`protocols/anti-capture-protocol.md`、`references/concept-cards/evidence-cost.md`、`references/concept-cards/power-closure.md`、`references/v2-source-spine.md`", "`public-power-governance-pack`、`judgment-responsibility-pack`、`framework-use-discipline-pack`"),
        ("框架边界 / 万能理论 / 专业替代 / 概念武器化", "`protocols/framework-boundary-protocol.md`、`worksheets/framework-boundary-check.md`、`references/framework-ontology-protection.md`、`templates/framework-boundary-output.md`", "`references/concept-cards/evidence-cost.md`、`protocols/anti-capture-protocol.md`", "`framework-use-discipline-pack`"),
        ("生命周期 / 阶段判断 / 长期演化过程", "`protocols/lifecycle-diagnosis-protocol.md`、`worksheets/lifecycle-stage-record.md`、`templates/lifecycle-output.md`", "`references/theory-backend-index.md`、`references/concept-cards/structure-process-group.md`、`references/v2-section-digest-index.md`", "`long-evolution-deep-pack`、`diagnosis-mainline-pack`"),
        ("递进模式 / 子锚点闭环 / 忙但没积累", "`protocols/progression-protocol.md`、`worksheets/sub-anchor-progression.md`、`templates/progression-output.md`", "`references/concept-cards/anchor-group.md`、`references/concept-cards/structure-process-group.md`", "`long-evolution-deep-pack`、`diagnosis-mainline-pack`"),
        ("势场 / 沉积基本盘 / 自主解离", "`protocols/field-dissociation-protocol.md`、`worksheets/field-dissociation-check.md`、`references/field-management-and-dissociation.md`、`templates/field-dissociation-output.md`", "`references/concept-cards/exit-transfer.md`、`references/concept-cards/repair-byproduct.md`", "`long-evolution-deep-pack`"),
        ("治理连续性 / 偿付约束 / 多中心治理 / 承接者生成", "`protocols/governance-continuity-protocol.md`、`worksheets/governance-continuity-check.md`、`templates/governance-continuity-output.md`", "`protocols/public-institution-protocol.md`、`references/theory-backend-index.md`", "`public-power-governance-pack`、`long-evolution-deep-pack`"),
        ("超大规模圈层 / 文明尺度 / 历史尺度压力测试", "`protocols/large-scale-stress-test-protocol.md`、`worksheets/large-scale-stress-test.md`、`templates/large-scale-stress-output.md`", "`references/theory-backend-index.md`、`references/concept-cards/scale-transfer.md`、`references/v2-source-spine.md`", "`long-evolution-deep-pack`、`public-power-governance-pack`、`framework-use-discipline-pack`"),
        ("对外表达 / 说人话 / 管理制度技术语境翻译", "`protocols/expression-translation-protocol.md`、`references/expression-translation-table.md`、`templates/expression-translation-output.md`、`templates/user-facing-language.md`", "对应诊断协议和概念卡", "`expression-article-pack`"),
        ("概念解释 / 思想解释 / 某概念怎么看", "`protocols/concept-explanation-protocol.md`、`references/concepts-minimal-set.md`、`references/v2-term-fidelity.md`、`templates/concept-explanation-output.md`", "与概念相关的概念卡；深度解释读 `v2-section-digest-index.md`", "`framework-use-discipline-pack`、`expression-article-pack`"),
        ("哲学 / 意义 / 第一因 / 生命是什么 / 虚无主义", "`protocols/concept-explanation-protocol.md`、`templates/concept-explanation-output.md`、`references/v2-term-fidelity.md`", "`concept-cards/scale-transfer.md`、`concept-cards/open-assertion.md`、`protocols/framework-boundary-protocol.md`", "`framework-use-discipline-pack`、`judgment-responsibility-pack`、`expression-article-pack`"),
        ("爱、牺牲、照护、公共承诺", "`references/concept-cards/love-open-action.md`、`references/love-as-open-action.md`", "`references/concept-cards/repair-byproduct.md`、`references/concept-cards/responsibility-chain.md`", "`intimate-love-care-pack`"),
        ("文明尺度 / 长期演化 / 制度生成 / 多中心治理 / 深层理论", "对应专项协议 + `references/theory-backend-index.md`", "`references/diagnostic-dimensions.md`、`references/diagnostic-toolbox-index.md`、`references/v2-coverage-map.md`、`references/v2-source-spine.md`", "`long-evolution-deep-pack`、`public-power-governance-pack`"),
    ]
    intro += "| 用户请求 | 必读 | 按需追加 | 连续联读包 |\n| --- | --- | --- | --- |\n"
    for row in rows:
        intro += "| " + " | ".join(row) + " |\n"
    intro += (
        "\n## 连续联读执行规则\n\n"
        "- 只要上表的“连续联读包”不是空，就先读 `references/continuity-bundles.md`，确认该包的同读材料和降档规则。\n"
        "- 深度、审计、高责任、公共制度、亲密关系、长期演化和文章输出场景，必须按需读 `references/v2-source-spine.md` 或 `references/v2-section-digest-index.md`，确认原文相邻章节。\n"
        "- 输出前使用 `worksheets/source-continuity-check.md`：若发现只读了单张概念卡，且本文件要求联读，必须补读或降档。\n"
        "- `templates/reasoning-outline-output.md` 中的“本次连续联读包”只列包名，不展开完整工作表。\n\n"
    )
    path.write_text(intro + high_risk, encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate v2 source continuity maps for CrossFrame.")
    parser.add_argument("--docx", default=r"D:\下载\跨尺度结构诊断框架v2.0.docx")
    parser.add_argument("--repo", default=".")
    args = parser.parse_args()

    repo = Path(args.repo).resolve()
    docx_path = Path(args.docx)
    write_source_files(repo, docx_path)
    update_read_routing(repo)
    print("generated v2 source spine, digest index, and read routing continuity table")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
