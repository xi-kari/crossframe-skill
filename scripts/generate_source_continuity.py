from __future__ import annotations

import argparse
import hashlib
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document


V5_DOCX = r"E:\世界模型\跨尺度结构诊断框架v5.0.docx"


@dataclass(frozen=True)
class Bundle:
    id: str
    name: str
    keywords: tuple[str, ...]
    reason: str
    trigger: str
    must_read: str
    failure: str
    downgrade: str
    output: str
    adjacent: tuple[str, ...]


V5_BUNDLES: tuple[Bundle, ...] = (
    Bundle(
        "v5-use-boundary-governance-pack",
        "使用边界与治理总闸包",
        ("使用边界", "使用准则", "输出边界", "适用性", "本体保护", "治理", "门禁", "总闸"),
        "把能否使用框架、能用到什么强度、是否需要暂停或降级放在所有诊断之前。",
        "用户要求 CrossFrame 直接定性、公开发布、替代专业判断、跨域迁移或把框架作为权威背书。",
        "使用边界、输出边界、适用性分级、本体保护、公开发布判断门禁。",
        "把框架当万能审判、专业替代、人格定性工具或组织背书材料。",
        "证据、对象、权力或用途任一不清时，只保留轻量观察或开放断言。",
        "先写可用范围和不可用范围，再写判断档位。",
        ("v5-diagnosis-admission-downgrade-exit-pack", "v5-framework-self-diagnosis-falsification-pack"),
    ),
    Bundle(
        "v5-diagnosis-admission-downgrade-exit-pack",
        "诊断准入、降级与退出包",
        ("诊断准入", "降级", "退出", "诊断可行性", "契约", "低条件", "行动上限"),
        "确保诊断不是无条件启动，而是在材料、用途、风险和退出空间满足最低条件后才进入。",
        "材料不足、风险紧急、用户要求快速判断、关系或组织没有正式复核条件。",
        "诊断准入、降级规则、退出条件、低条件试探行动、行动上限。",
        "证据不足仍维持强判断，或风险高却不给退出路径。",
        "降到低风险、可撤回、可观察的小动作，并声明暂停判断条件。",
        "输出准入结论、降级理由、行动上限和撤回条件。",
        ("v5-low-power-protection-pack", "v5-evidence-downgrade-action-ceiling-pack"),
    ),
    Bundle(
        "v5-low-power-protection-pack",
        "低权力主体保护包",
        ("低权力", "弱信号", "反报复", "无法退出", "保护", "安全", "匿名", "低暴露"),
        "先保护承受成本更高、表达更危险、退出空间更小的主体，防止诊断扩大伤害。",
        "亲密关系、家庭、小团队、平台治理、公共制度、投诉申诉、无法退出主体。",
        "低权力主体、弱信号保护、反报复、低暴露记录、退出转移。",
        "把修复责任压回受伤者，或让弱信号暴露在高权力对象面前。",
        "不能保护时不得公开强判断，只能给安全边界和外部承接建议。",
        "输出必须先写安全、暴露、退出和保护边界。",
        ("v5-love-trapped-trauma-pack", "v5-public-power-institution-pack"),
    ),
    Bundle(
        "v5-concept-weaponization-dogma-pack",
        "概念武器化与教条化包",
        ("概念武器化", "教条", "人格审判", "术语", "概念有效性", "万能", "误用"),
        "阻止概念从观察工具变成处置武器、身份标签或组织话术。",
        "用户要求贴标签、定性人品、用概念替代证据或把术语直接写成结论。",
        "概念有效性分级、概念武器化防范、框架教条化防范、表达翻译规则。",
        "用术语制造结论，或把低级别概念探索升级为强处置依据。",
        "概念只保留为候选语言，不能承担强判断或处置作用。",
        "输出必须把术语翻译成现实行为、证据和可撤回条件。",
        ("v5-core-concept-integrity-pack", "v5-domain-translation-normative-source-pack"),
    ),
    Bundle(
        "v5-cross-scale-context-translation-pack",
        "跨尺度与语境转译包",
        ("跨尺度", "尺度", "语境", "转译", "前概念闸", "隐喻", "迁移", "层级"),
        "防止从个体、组织、制度、历史之间直接偷换尺度或用隐喻冒充证明。",
        "问题涉及升维解释、跨域类比、制度化语言、技术治理或普通人表达。",
        "跨尺度迁移闸、前概念闸、语境翻译、隐喻漂移控制。",
        "用一个尺度的事实直接证明另一个尺度的因果或责任。",
        "不能通过前概念闸时，只能说明相似性，不能做因果迁移。",
        "输出必须声明对象、层级、时间窗口和转译损失。",
        ("v5-domain-translation-normative-source-pack", "v5-state-coordinate-lifecycle-pack"),
    ),
    Bundle(
        "v5-seven-gates-diagnosis-pack",
        "七闸诊断包",
        ("七闸", "五闸", "诊断主流程", "对象闸", "证据闸", "尺度闸", "责任闸", "观测闸", "权力闸", "行动闸"),
        "把 v5 诊断从五闸升级为七闸复核，防止证据、权力和行动上限被跳过。",
        "所有完整诊断、深度分析、高责任判断、公共制度、组织修复和文章底稿。",
        "对象、证据、尺度、责任、观测、权力、行动七个闸口。",
        "只套机制名，不交代对象、证据、权力结构、观测影响或行动上限。",
        "七闸任一缺失时降到开放断言或低条件行动。",
        "输出提纲必须列出七闸结论或说明哪个闸导致降级。",
        ("v5-source-evidence-separation-pack", "v5-evidence-downgrade-action-ceiling-pack"),
    ),
    Bundle(
        "v5-source-evidence-separation-pack",
        "来源、证据与判断分离包",
        ("来源", "证据", "判断", "材料", "事实", "自评", "缺失材料", "证据成本"),
        "把材料来源、可核验证据、解释候选和判断档位分开，防止文本漂亮就被当成现实证明。",
        "有报告、自评、聊天记录、媒体材料、AI 生成材料、申诉材料或不完整证据。",
        "来源栏、证据栏、判断栏、缺失材料清单、证据成本。",
        "把来源权威、格式完整或叙述顺滑误当作高成本证据。",
        "证据成本不够时降低判断档位，并列出待补材料。",
        "输出必须显式区分事实、解释、机制候选和判断档位。",
        ("v5-ai-process-artifact-boundary-pack", "v5-evidence-downgrade-action-ceiling-pack"),
    ),
    Bundle(
        "v5-evidence-downgrade-action-ceiling-pack",
        "证据降级与行动上限包",
        ("证据降级", "行动上限", "判断档位", "开放断言", "低条件", "撤回", "反向条件"),
        "让证据不足时的判断自动降低行动强度，避免低证据支撑高处置。",
        "任何涉及名誉、资源、处分、组织处置、公共记忆、公开评论或关系退出的判断。",
        "证据降级、判断档位、行动上限、撤回条件、反向条件。",
        "用开放断言支撑高责任行动，或给出无法撤回的结论。",
        "证据不够时只能给观察、开放断言、低条件行动或补证清单。",
        "输出必须写清判断能做什么、不能做什么、何时撤回。",
        ("v5-open-assertion-proposition-pack", "v5-strong-judgment-eight-pack"),
    ),
    Bundle(
        "v5-observation-reflexivity-release-pack",
        "观测反身性与收束包",
        ("观测", "反身", "递归", "收束", "高反身性", "表演", "反制", "熵增"),
        "处理被诊断、命名、公开或追踪后对象会改变行为的场景，并规定何时停止递归观察。",
        "对象会表演、反制、改证据、改边界，或用户要求持续追踪反应。",
        "观测影响登记、高反身性博弈、观测递归扩张与收束、阶段 6 熵增边界。",
        "无限追踪反应，或把对象被观察后的表演当作原始证据。",
        "第三层后无新增结构变量或风险增加时必须收束。",
        "输出必须区分原始状态、观测后反应和收束理由。",
        ("v5-state-coordinate-lifecycle-pack", "v5-media-platform-crisis-pack"),
    ),
    Bundle(
        "v5-ai-process-artifact-boundary-pack",
        "AI 过程性产物边界包",
        ("AI", "人工智能", "过程性产物", "合规", "报告", "模板", "自评", "L1-L3"),
        "防止 AI 报告、合规模板、整理稿或过程性产物被当作现实已经被验证。",
        "AI 报告、合规材料、机构自评、审计稿、流程图、自动生成证据或模型诊断。",
        "AI 输出强制声明、缺失材料清单、过程性产物不得充当现实证明、L1-L3 模板。",
        "把 AI 产物的完整性误认为现实处置、反报复保护或独立核验已经发生。",
        "AI 产物只能作为线索整理，不能独立支撑强判断。",
        "输出必须声明 AI 材料边界、缺失现实材料和可核验下一步。",
        ("v5-source-evidence-separation-pack", "v5-concept-weaponization-dogma-pack"),
    ),
    Bundle(
        "v5-strong-judgment-eight-pack",
        "强判断八件套包",
        ("强判断", "八件套", "发布门禁", "资格", "名誉", "资源", "权利", "处分"),
        "高责任判断必须先通过强判断八件套，避免把结构判断变成不可申诉的处置。",
        "涉及资格、名誉、权利、资源、处分、公共记忆、公开点名或重大组织行动。",
        "强判断十问、强判断八件套、发布门禁、反证入口、申诉和修复路径。",
        "缺少反证、申诉、撤回、补偿或独立复核却公开强判断。",
        "八件套不全时不能强判断，只能开放断言或命题验证待办。",
        "输出必须列出强判断是否成立、缺项和撤回路径。",
        ("v5-open-assertion-proposition-pack", "v5-low-power-protection-pack"),
    ),
    Bundle(
        "v5-open-assertion-proposition-pack",
        "开放断言与命题验证包",
        ("开放断言", "命题验证", "前瞻登记", "反证", "申诉", "可撤回", "判断勇气"),
        "保留必要判断勇气，同时让判断停留在可撤回、可验证、可申诉的层级。",
        "证据不足但风险不能无视、用户要求判断、需要提出命题或预测。",
        "开放断言记录表、命题验证表、前瞻登记、反证入口、撤回条件。",
        "用开放断言绕过证据责任，或把命题验证写成已经证明。",
        "未完成验证前不得进入强处置或公共记忆。",
        "输出必须写命题、证据状态、反向条件和验证窗口。",
        ("v5-evidence-downgrade-action-ceiling-pack", "v5-strong-judgment-eight-pack"),
    ),
    Bundle(
        "v5-core-concept-integrity-pack",
        "核心概念完整性包",
        ("核心概念", "概念层级", "锚点", "动力组", "结构组", "过程组", "承接", "回流", "完整性"),
        "保持 v5 核心概念在锚点、动力、结构、过程之间的连续关系，避免只抽一个词使用。",
        "输出要使用承接、回流、锚点、结构负荷、修复副产品、责任链等高风险概念。",
        "核心概念最小集、概念层级、锚点组、动力组、结构组、过程组、概念融合规则。",
        "只读单张概念卡就让概念承担完整诊断作用。",
        "未读完整相邻概念时，概念只能作为解释线索。",
        "输出必须把概念落回现实行为、关系、成本和反馈。",
        ("v5-anchor-dynamics-structure-process-pack", "v5-concept-weaponization-dogma-pack"),
    ),
    Bundle(
        "v5-anchor-dynamics-structure-process-pack",
        "锚点、动力、结构与过程包",
        ("锚点", "动力", "结构", "过程", "行动承接", "支撑通道", "条件场", "结构负荷"),
        "把对象的稳定点、动力来源、结构负荷和过程写回放在同一条诊断链上。",
        "解释一个系统为什么能启动、卡住、转译失败、责任断裂或忙而无积累。",
        "锚点组、动力组、结构组、过程组、行动承接和修复副产品。",
        "只谈价值或意愿，不追踪支撑通道、结构负荷和反馈写回。",
        "缺少过程证据时不得下结构定论，只能列机制候选。",
        "输出必须写主锚点、动力通道、结构负荷和过程写回状态。",
        ("v5-core-concept-integrity-pack", "v5-long-evolution-progression-field-pack"),
    ),
    Bundle(
        "v5-root-assumptions-meta-rules-pack",
        "根假设与元规则包",
        ("根假设", "元规则", "核心推论", "暂停使用", "反例", "假设层", "边界元假设"),
        "让框架底层假设可声明、可降级、可暂停、可接受反例，而不是用定义自保。",
        "用户问框架是否成立、概念是否自洽、理论后台、第一因、意义或抽象命题。",
        "根假设层级、元规则、核心推论、开放行动与非闭合、根假设使用纪律。",
        "用框架自己的概念证明框架永远正确。",
        "反例触发时暂停对应根假设，不强行补丁化解释。",
        "输出必须声明使用了哪个根假设、可被什么反例暂停。",
        ("v5-framework-self-diagnosis-falsification-pack", "v5-domain-translation-normative-source-pack"),
    ),
    Bundle(
        "v5-state-coordinate-lifecycle-pack",
        "状态坐标与生命周期包",
        ("状态坐标", "阶段", "生命周期", "阶段0", "阶段6", "时间窗口", "局部状态", "全周期"),
        "把阶段 0-6 改写成局部状态坐标，禁止写成线性宿命或所有对象必经路径。",
        "生命周期判断、阶段判断、发展阶段、成熟/衰退、长期演化或阶段 6。",
        "状态坐标、对象、层级、时间窗口、并行子系统、反向条件、暂停条件、撤回路径。",
        "把阶段写成命运铁轨，或忽略对象和时间窗口。",
        "无法声明坐标条件时，不得给阶段结论。",
        "输出必须写对象、层级、时间窗口、反向条件、暂停条件和撤回路径。",
        ("v5-observation-reflexivity-release-pack", "v5-long-evolution-progression-field-pack"),
    ),
    Bundle(
        "v5-long-evolution-progression-field-pack",
        "长期演化、递进与势场包",
        ("长期", "递进", "势场", "自主解离", "基本盘", "沉积", "预警", "多中心", "代际"),
        "处理长期累积、子锚点递进、正负势场、沉积基本盘、自主解离与治理连续性。",
        "战略推进、长期修复、忙但没有积累、系统停滞、组织演化、文明尺度压力测试。",
        "递进模式、双向势场、自主解离、调节预警偿付、多中心治理、代际承接。",
        "把势场当氛围，把长期演化写成单因果线。",
        "缺少时间序列和反馈写回时，只能给观察框架。",
        "输出必须写递进链、势场方向、承接者和反馈写回。",
        ("v5-governance-continuity-multicenter-pack", "v5-state-coordinate-lifecycle-pack"),
    ),
    Bundle(
        "v5-governance-continuity-multicenter-pack",
        "治理连续性与多中心包",
        ("治理连续性", "多中心", "调节", "预警", "偿付", "承接者", "代际", "嵌套治理"),
        "检查一个系统是否有持续调节、预警、偿付、承接者再生产和多中心纠偏能力。",
        "组织修复、制度治理、长期项目、公共承诺、代际承接或平台治理。",
        "调节-控制、监测-深时间预警、承诺-偿付、多中心-嵌套治理、代际承接。",
        "把一次性修复、口号承诺或单中心控制当作治理连续性。",
        "没有偿付和承接者生成时，只能说短期管理，不能说长期治理。",
        "输出必须写调节、预警、偿付、多中心和承接者状态。",
        ("v5-public-power-institution-pack", "v5-long-evolution-progression-field-pack"),
    ),
    Bundle(
        "v5-action-healing-transfer-pack",
        "行动、疗愈与转移包",
        ("行动", "疗愈", "转移", "修复", "退出", "重建", "操作层", "爱诊断红线"),
        "把诊断落到低风险行动、疗愈边界、退出转移和修复写回，而不是停留在解释。",
        "用户要下一步、修复方案、疗愈、退出、关系/组织重建或低条件试探行动。",
        "操作层通则、疗愈方案、退出转移、爱诊断红线、修复副产品。",
        "把疗愈替代医疗法律安全处置，或把修复责任压给受害者。",
        "高风险场景只给边界和转介，不给伪专业处置。",
        "输出必须写下一步行动、风险边界、承接者和停止条件。",
        ("v5-responsibility-intervention-separation-pack", "v5-love-trapped-trauma-pack"),
    ),
    Bundle(
        "v5-responsibility-intervention-separation-pack",
        "责任链与干预边界分离包",
        ("责任", "干预", "责任链", "承接者", "成本", "中层耗竭", "高责任", "操作准则"),
        "分清谁造成问题、谁有条件改变、谁在承接成本、诊断者能干预到哪里。",
        "组织复盘、公共评论、亲密关系、团队修复、问责和行动建议。",
        "责任链、行动承接、结构负荷、干预边界、成本分配和修复写回。",
        "把责任平均化、抽象化，或让无权者承担修复成本。",
        "责任链不清时只给问题拆分和待核验清单。",
        "输出必须写责任位置、可干预位置和不可替代的外部处置。",
        ("v5-low-power-protection-pack", "v5-action-healing-transfer-pack"),
    ),
    Bundle(
        "v5-love-trapped-trauma-pack",
        "爱、无法退出与复杂创伤包",
        ("爱", "无法退出", "复杂创伤", "健康基准", "照护", "创伤", "开放性承担", "保护优先级"),
        "处理爱、照护、无法退出和复杂创伤时先保护痛苦、安全、主体性和退出空间。",
        "亲密关系、家庭、照护、创伤性生存策略、无健康基准、爱被要求或道德化。",
        "无法退出主体保护、复杂创伤、无健康基准、开放性承担行动、爱诊断红线。",
        "把爱写成忍耐命令，或要求无法退出者完成理想退出。",
        "不能保障安全时，不做关系修复建议，只给保护和外部承接。",
        "输出必须先接住处境，再写边界、保护、承接和可撤回行动。",
        ("v5-low-power-protection-pack", "v5-action-healing-transfer-pack"),
    ),
    Bundle(
        "v5-public-power-institution-pack",
        "公共权力与制度包",
        ("公共", "制度", "平台", "权力", "公共承诺", "程序", "高权力密度", "治理"),
        "公共权力和制度场景必须处理程序有效性、反报复、证据降级和行动上限。",
        "平台治理、政策、机构责任、公共争议、公共评论、申诉、合规和制度修复。",
        "公共承诺、高权力密度、程序有效性、制度治理、反报复和公共记忆。",
        "把程序外观当有效程序，或让公共评论越过证据边界。",
        "程序无效或反报复不足时不能给最终性公共强判断。",
        "输出必须写制度责任、程序缺口、证据边界和低权力保护。",
        ("v5-media-platform-crisis-pack", "v5-governance-continuity-multicenter-pack"),
    ),
    Bundle(
        "v5-domain-translation-normative-source-pack",
        "领域转译、规范前提与来源透明包",
        ("领域", "翻译", "规范性", "来源", "知识谱系", "表达", "隐喻", "价值边界"),
        "公开框架的价值边界、知识来源和转译损失，避免伪装成价值中立或吞并外部理论。",
        "对外文章、读书互读、理论解释、跨学科转译、概念阐释和公共评论。",
        "规范性前提、知识谱系、来源透明、隐喻漂移、对外表达规则。",
        "把价值选择伪装成纯经验结论，或用隐喻证明因果。",
        "来源和规范前提不清时，降低为观点阐释而非事实诊断。",
        "输出必须声明价值边界、来源关系和可争议处。",
        ("v5-cross-scale-context-translation-pack", "v5-root-assumptions-meta-rules-pack"),
    ),
    Bundle(
        "v5-media-platform-crisis-pack",
        "媒体、平台与危机包",
        ("媒体", "平台", "舆论", "危机", "公共记忆", "发布", "传播", "申诉"),
        "处理传播放大、平台权力、公共记忆和危机响应中的证据、反报复与发布责任。",
        "媒体事件、平台处罚、公共争议、危机公关、公开发文、舆论反转。",
        "公共权力、证据降级、弱信号保护、发布门禁、公共记忆和申诉入口。",
        "用传播热度替代证据，或把未验证判断写入公共记忆。",
        "公开材料未过门禁时，只能写边界化评论和待核验问题。",
        "输出必须写传播风险、证据状态、发布条件和撤回机制。",
        ("v5-public-power-institution-pack", "v5-strong-judgment-eight-pack"),
    ),
    Bundle(
        "v5-framework-self-diagnosis-falsification-pack",
        "框架自诊、证伪与回滚包",
        ("框架自诊", "证伪", "版本治理", "反例", "案例库", "幸存者偏差", "良性消亡", "回滚"),
        "让 CrossFrame 本身也接受诊断、反例、版本治理、失败样本和必要退场。",
        "用户要求评估框架、比较 3.0/5.0、检查输出是否执行框架或质疑框架有效性。",
        "自诊对象、版本治理、根假设证伪、案例库偏差、判断追踪、良性消亡。",
        "用框架语言自我证明，或为了保住框架无限加补丁。",
        "反例成立时暂停相关模块，保留替代框架接口。",
        "输出必须写框架适用对象、反例入口、回滚条件和替代接口。",
        ("v5-root-assumptions-meta-rules-pack", "v5-toolization-accessibility-release-pack"),
    ),
    Bundle(
        "v5-toolization-accessibility-release-pack",
        "工具化、可及性与发布包",
        ("工具化", "可及性", "使用门槛", "商业化", "认证", "分裂", "公开使用", "发布"),
        "控制框架被课程、咨询、AI 工具或组织软件使用时产生的新解释权不平等。",
        "技能封装、工具化、公开发布、商业化、培训认证、自动化审查或大规模部署。",
        "使用门槛债、可及性审计、工具化红线、分裂协议、公开使用边界。",
        "让框架变成专家垄断、商业交付压力或新的权力工具。",
        "无法保证可及性和复核时，只能内部试用或降级发布。",
        "输出必须写使用门槛、复核责任、公开边界和回滚方式。",
        ("v5-use-boundary-governance-pack", "v5-framework-self-diagnosis-falsification-pack"),
    ),
)


REQUIRED_WITH: dict[str, tuple[str, ...]] = {
    "v5-use-boundary-governance-pack": (
        "v5-diagnosis-admission-downgrade-exit-pack",
        "v5-concept-weaponization-dogma-pack",
    ),
    "v5-diagnosis-admission-downgrade-exit-pack": (
        "v5-seven-gates-diagnosis-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
    ),
    "v5-low-power-protection-pack": (
        "v5-source-evidence-separation-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
    ),
    "v5-concept-weaponization-dogma-pack": (
        "v5-core-concept-integrity-pack",
        "v5-use-boundary-governance-pack",
    ),
    "v5-cross-scale-context-translation-pack": (
        "v5-domain-translation-normative-source-pack",
    ),
    "v5-seven-gates-diagnosis-pack": (
        "v5-source-evidence-separation-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
    ),
    "v5-source-evidence-separation-pack": (
        "v5-evidence-downgrade-action-ceiling-pack",
    ),
    "v5-evidence-downgrade-action-ceiling-pack": (),
    "v5-observation-reflexivity-release-pack": (
        "v5-source-evidence-separation-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
    ),
    "v5-ai-process-artifact-boundary-pack": (
        "v5-source-evidence-separation-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
        "v5-concept-weaponization-dogma-pack",
    ),
    "v5-strong-judgment-eight-pack": (
        "v5-open-assertion-proposition-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
        "v5-low-power-protection-pack",
        "v5-seven-gates-diagnosis-pack",
    ),
    "v5-open-assertion-proposition-pack": (
        "v5-source-evidence-separation-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
    ),
    "v5-core-concept-integrity-pack": (
        "v5-anchor-dynamics-structure-process-pack",
    ),
    "v5-anchor-dynamics-structure-process-pack": (
        "v5-core-concept-integrity-pack",
    ),
    "v5-root-assumptions-meta-rules-pack": (
        "v5-domain-translation-normative-source-pack",
    ),
    "v5-state-coordinate-lifecycle-pack": (
        "v5-observation-reflexivity-release-pack",
    ),
    "v5-long-evolution-progression-field-pack": (
        "v5-state-coordinate-lifecycle-pack",
        "v5-governance-continuity-multicenter-pack",
    ),
    "v5-governance-continuity-multicenter-pack": (
        "v5-public-power-institution-pack",
        "v5-responsibility-intervention-separation-pack",
    ),
    "v5-action-healing-transfer-pack": (
        "v5-responsibility-intervention-separation-pack",
        "v5-low-power-protection-pack",
    ),
    "v5-responsibility-intervention-separation-pack": (
        "v5-source-evidence-separation-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
    ),
    "v5-love-trapped-trauma-pack": (
        "v5-low-power-protection-pack",
        "v5-action-healing-transfer-pack",
    ),
    "v5-public-power-institution-pack": (
        "v5-low-power-protection-pack",
        "v5-source-evidence-separation-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
    ),
    "v5-domain-translation-normative-source-pack": (
        "v5-use-boundary-governance-pack",
        "v5-concept-weaponization-dogma-pack",
    ),
    "v5-media-platform-crisis-pack": (
        "v5-public-power-institution-pack",
        "v5-evidence-downgrade-action-ceiling-pack",
        "v5-strong-judgment-eight-pack",
    ),
    "v5-framework-self-diagnosis-falsification-pack": (
        "v5-root-assumptions-meta-rules-pack",
        "v5-use-boundary-governance-pack",
    ),
    "v5-toolization-accessibility-release-pack": (
        "v5-use-boundary-governance-pack",
        "v5-framework-self-diagnosis-falsification-pack",
        "v5-low-power-protection-pack",
    ),
}


PART_FALLBACK = {
    "第一部分": ("v5-use-boundary-governance-pack", "v5-diagnosis-admission-downgrade-exit-pack"),
    "第二部分": ("v5-use-boundary-governance-pack", "v5-concept-weaponization-dogma-pack"),
    "第三部分": ("v5-core-concept-integrity-pack", "v5-state-coordinate-lifecycle-pack"),
    "第四部分": ("v5-seven-gates-diagnosis-pack", "v5-source-evidence-separation-pack"),
    "第五部分": ("v5-action-healing-transfer-pack", "v5-responsibility-intervention-separation-pack"),
    "第六部分": ("v5-public-power-institution-pack", "v5-media-platform-crisis-pack"),
    "第七部分": ("v5-framework-self-diagnosis-falsification-pack", "v5-toolization-accessibility-release-pack"),
}


ROUTES = (
    ("快速诊断", "protocols/diagnosis-protocol.md; worksheets/intake-worksheet.md; worksheets/seven-gates-worksheet.md", ("v5-diagnosis-admission-downgrade-exit-pack", "v5-seven-gates-diagnosis-pack")),
    ("完整诊断 / 审计 / 深度分析", "protocols/diagnosis-protocol.md; references/v5-section-digest-index.md; references/v5-term-fidelity.md", ("v5-seven-gates-diagnosis-pack", "v5-source-evidence-separation-pack", "v5-core-concept-integrity-pack")),
    ("开放断言 / 命题验证", "protocols/open-assertion-protocol.md; protocols/proposition-verification-protocol.md", ("v5-open-assertion-proposition-pack", "v5-evidence-downgrade-action-ceiling-pack")),
    ("公开强判断 / 资格名誉资源权利", "protocols/proposition-verification-protocol.md; protocols/anti-capture-protocol.md", ("v5-strong-judgment-eight-pack", "v5-low-power-protection-pack", "v5-evidence-downgrade-action-ceiling-pack")),
    ("AI 报告 / 合规材料 / 自评", "references/concept-cards/malicious-compliance-ai-validation.md; references/concept-cards/evidence-cost.md", ("v5-ai-process-artifact-boundary-pack", "v5-source-evidence-separation-pack")),
    ("生命周期 / 阶段判断 / 长期演化", "protocols/lifecycle-diagnosis-protocol.md; protocols/progression-protocol.md", ("v5-state-coordinate-lifecycle-pack", "v5-long-evolution-progression-field-pack")),
    ("组织复盘 / 修复", "protocols/governance-continuity-protocol.md; protocols/healing-transfer-protocol.md", ("v5-governance-continuity-multicenter-pack", "v5-responsibility-intervention-separation-pack", "v5-action-healing-transfer-pack")),
    ("公共制度 / 平台治理 / 公共评论", "protocols/public-institution-protocol.md; protocols/anti-capture-protocol.md", ("v5-public-power-institution-pack", "v5-media-platform-crisis-pack", "v5-low-power-protection-pack")),
    ("亲密关系 / 无法退出 / 创伤", "protocols/intimate-relationship-protocol.md; protocols/healing-transfer-protocol.md", ("v5-love-trapped-trauma-pack", "v5-low-power-protection-pack", "v5-action-healing-transfer-pack")),
    ("文章 / 评论 / 可读输出", "../crossframe-essay/SKILL.md; templates/user-facing-language.md", ("v5-domain-translation-normative-source-pack", "v5-seven-gates-diagnosis-pack")),
    ("框架自诊 / 证伪 / 版本治理", "protocols/framework-boundary-protocol.md; references/framework-ontology-protection.md", ("v5-framework-self-diagnosis-falsification-pack", "v5-root-assumptions-meta-rules-pack")),
    ("工具化 / 公开发布 / 技能封装", "references/framework-ontology-protection.md; worksheets/source-continuity-check.md", ("v5-toolization-accessibility-release-pack", "v5-use-boundary-governance-pack")),
)


def normalize(text: str) -> str:
    return " ".join(text.split())


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def heading_level(style_name: str) -> int:
    if style_name.startswith("Heading"):
        match = re.search(r"(\d+)", style_name)
        return int(match.group(1)) if match else 1
    if style_name.startswith("标题"):
        match = re.search(r"(\d+)", style_name)
        if match:
            value = match.group(1)
            if value.startswith("3"):
                return 3
            if value.startswith("2"):
                return 2
        return 2
    return 9


def extract_tables(doc: Document) -> list[dict[str, object]]:
    tables: list[dict[str, object]] = []
    for index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows[:4]:
            cells = [normalize(cell.text) for cell in row.cells[:6]]
            if any(cells):
                rows.append(cells)
        tables.append({"id": f"V5-T{index:03d}", "rows": rows})
    return tables


def extract_nodes(path: Path) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    doc = Document(str(path))
    paragraphs = [
        {"index": i, "style": para.style.name, "text": normalize(para.text)}
        for i, para in enumerate(doc.paragraphs)
    ]
    start = None
    for item in paragraphs:
        if item["text"] == "第一部分：极简导读" and is_heading_style(str(item["style"])):
            start = int(item["index"])
            break
    if start is None:
        raise SystemExit("cannot find styled v5 body start: 第一部分：极简导读")

    heading_positions = [
        i
        for i, item in enumerate(paragraphs)
        if int(item["index"]) >= start and item["text"] and is_heading_style(str(item["style"]))
    ]

    nodes: list[dict[str, object]] = []
    current_part = "第一部分：极简导读"
    for n, pos in enumerate(heading_positions, start=1):
        item = paragraphs[pos]
        next_pos = heading_positions[n] if n < len(heading_positions) else len(paragraphs)
        title = str(item["text"])
        if str(item["style"]).startswith("Heading"):
            current_part = title
        body_items = [p for p in paragraphs[pos + 1 : next_pos] if p["text"]]
        body_text = "\n".join(str(p["text"]) for p in body_items)
        digest = make_digest(body_text)
        node_text = f"{title}\n{digest}"
        bundle_ids = bundles_for_node(node_text, current_part)
        prev_id = f"V5-H{n - 1:03d}" if n > 1 else ""
        next_id = f"V5-H{n + 1:03d}" if n < len(heading_positions) else ""
        end_para = int(paragraphs[next_pos - 1]["index"]) if next_pos > pos else int(item["index"])
        nodes.append(
            {
                "id": f"V5-H{n:03d}",
                "title": title,
                "style": str(item["style"]),
                "level": heading_level(str(item["style"])),
                "part": current_part,
                "start": int(item["index"]),
                "end": end_para,
                "body_chars": len(body_text),
                "digest": digest or "本节是标题型节点，正文内容与相邻节点共同承载。",
                "prev": prev_id,
                "next": next_id,
                "bundles": bundle_ids,
            }
        )
    return nodes, extract_tables(doc)


def is_heading_style(style_name: str) -> bool:
    return style_name.startswith("Heading") or style_name.startswith("标题")


def make_digest(body_text: str) -> str:
    clean = normalize(body_text.replace("|", " "))
    if not clean:
        return ""
    pieces = re.split(r"(?<=[。！？；])", clean)
    digest = "".join(piece for piece in pieces[:3]).strip()
    if len(digest) < 30:
        digest = clean[:260]
    return digest[:360]


def bundles_for_node(text: str, part: str) -> list[str]:
    result: list[str] = []
    for bundle in V5_BUNDLES:
        if any(keyword in text for keyword in bundle.keywords):
            result.append(bundle.id)
    for key, ids in PART_FALLBACK.items():
        if key in part:
            result.extend(ids)
            break
    seen: set[str] = set()
    unique = [bundle_id for bundle_id in result if not (bundle_id in seen or seen.add(bundle_id))]
    return unique[:5]


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8", newline="\n")


def bundle_by_id(bundle_id: str) -> Bundle:
    for bundle in V5_BUNDLES:
        if bundle.id == bundle_id:
            return bundle
    raise KeyError(bundle_id)


def required_closure(bundle_ids: list[str] | tuple[str, ...]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set(bundle_ids)
    stack = list(bundle_ids)
    while stack:
        current = stack.pop(0)
        for required in REQUIRED_WITH.get(current, ()):
            if required not in seen:
                seen.add(required)
                result.append(required)
                stack.append(required)
    return result


def linked_bundles(ids: list[str]) -> str:
    return ", ".join(f"`{bundle_id}`" for bundle_id in ids) or "无"


def node_rows(nodes: list[dict[str, object]], bundle_id: str | None = None, limit: int | None = None) -> list[dict[str, object]]:
    selected = [node for node in nodes if bundle_id is None or bundle_id in node["bundles"]]
    return selected[:limit] if limit else selected


def render_source_spine(nodes: list[dict[str, object]], tables: list[dict[str, object]], source_docx: Path) -> str:
    lines = [
        "# CrossFrame v5 Source Spine",
        "",
        f"- 权威源：`{source_docx}`",
        f"- SHA256：`{sha256(source_docx)}`",
        f"- 标题节点：{len(nodes)}",
        f"- 表格数量：{len(tables)}",
        "- 抽取规则：从样式化 `第一部分：极简导读` 开始，读取 Word `Heading*` 与中文 `标题*` 样式；目录区重复标题不计入正文节点。",
        "- 用途：恢复 v5 原文顺序、段落范围、相邻关系与默认连读包；不要把本文件当正文替代品。",
        "",
        "## 标题顺序与相邻关系",
        "",
        "| ID | 层级 | 部分 | 段落范围 | 标题 | 前后 | 默认连读包 |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ]
    for node in nodes:
        prev_next = " / ".join(part for part in [str(node["prev"]), str(node["next"])] if part) or "-"
        lines.append(
            f"| `{node['id']}` | {node['level']} | {node['part']} | {node['start']}-{node['end']} | {node['title']} | {prev_next} | {linked_bundles(list(node['bundles']))} |"
        )
    lines.extend(["", "## 表格索引", "", "| ID | 抽取预览 |", "| --- | --- |"])
    for table in tables:
        preview = "；".join(" / ".join(cell for cell in row if cell) for row in table["rows"])
        lines.append(f"| `{table['id']}` | {preview[:260]} |")
    return "\n".join(lines)


def render_digest(nodes: list[dict[str, object]], source_docx: Path) -> str:
    lines = [
        "# CrossFrame v5 Section Digest Index",
        "",
        f"权威源：`{source_docx}`",
        "",
        "本文件逐标题节点保存 v5 的保真摘要、段落范围、相邻节点和默认连读包。摘要是研究笔记，不替代原文。",
        "",
    ]
    for node in nodes:
        lines.extend(
            [
                f"## `{node['id']}` {node['title']}",
                "",
                f"- 部分：{node['part']}",
                f"- 段落范围：{node['start']}-{node['end']}；正文字符：{node['body_chars']}",
                f"- 相邻节点：前 `{node['prev'] or '无'}`；后 `{node['next'] or '无'}`",
                f"- 默认连读包：{linked_bundles(list(node['bundles']))}",
                f"- 保真摘要：{node['digest']}",
                "- 不可误读：不能脱离相邻标题、七闸、证据降级和行动上限单独抽用。",
                "",
            ]
        )
    return "\n".join(lines)


def render_coverage(nodes: list[dict[str, object]], tables: list[dict[str, object]]) -> str:
    lines = [
        "# CrossFrame v5 Coverage Map",
        "",
        "本文件说明 v5 DOCX 的标题节点、连读包、协议和工作表之间的覆盖关系。",
        "",
        "## 覆盖总览",
        "",
        f"- v5 标题节点：{len(nodes)}",
        f"- v5 表格：{len(tables)}",
        f"- v5 连读包：{len(V5_BUNDLES)}",
        "- `continuity-bundles.md` 只做索引；每个连读包正文放在 `references/continuity-bundles/v5/`。",
        "",
        "## 连读包覆盖",
        "",
        "| 连读包 | 命中标题数 | 代表标题 | 主要协议/工作表 |",
        "| --- | ---: | --- | --- |",
    ]
    route_lookup = {bundle_id: [] for bundle_id in [b.id for b in V5_BUNDLES]}
    for route_name, materials, bundle_ids in ROUTES:
        for bundle_id in bundle_ids:
            route_lookup[bundle_id].append(f"{route_name}: {materials}")
    for bundle in V5_BUNDLES:
        selected = node_rows(nodes, bundle.id)
        samples = "；".join(str(node["title"]) for node in selected[:5]) or "按触发场景读取"
        materials = "<br>".join(route_lookup.get(bundle.id) or ["按 read-routing-map.md 场景读取"])
        lines.append(f"| `{bundle.id}` | {len(selected)} | {samples} | {materials} |")
    lines.extend(["", "## 技能覆盖", "", "| Skill | v5 读取责任 |", "| --- | --- |"])
    skill_rows = (
        ("crossframe", "核心七闸、v5 source modules、26 个连读包、概念保真和诊断协议。"),
        ("crossframe-suite", "只做模式/角色选择和路由；进入成文前交给 essay 的文章类型选择器；v5 包由目标 skill 读取。"),
        ("crossframe-essay", "保留文章类型选择器和 50 个写作技法；成文前读取 v5 路由包和结构洞察底稿。"),
        ("crossframe-review", "隐藏质量闸；检查七闸、八件套、低权力保护、证据降级和正文未被吞掉。"),
        ("crossframe-public / org / dialogue / debate / notebook / casebook / teach", "专项 skill 按 read-routing-map.md 读取对应 v5 连读包，不独立改写权威源。"),
    )
    for skill, responsibility in skill_rows:
        lines.append(f"| `{skill}` | {responsibility} |")
    return "\n".join(lines)


def render_term_fidelity(nodes: list[dict[str, object]]) -> str:
    terms: list[str] = []
    for bundle in V5_BUNDLES:
        terms.extend(bundle.keywords[:4])
    for extra in ("七闸", "强判断八件套", "局部状态坐标", "过程性产物", "开放性承担行动", "低条件试探行动"):
        terms.append(extra)
    seen: set[str] = set()
    terms = [term for term in terms if term and not (term in seen or seen.add(term))]
    lines = [
        "# CrossFrame v5 Term Fidelity",
        "",
        "本文件是 v5 术语保真表。使用任何术语承担判断作用前，必须同时检查其相邻连读包和现实表达。",
        "",
        "| 术语 | 保真要求 | 常见误读 | 默认连读包 |",
        "| --- | --- | --- | --- |",
    ]
    for term in terms:
        matched = [bundle.id for bundle in V5_BUNDLES if term in bundle.keywords][:3]
        if not matched:
            matched = [node["bundles"][0] for node in nodes if term in str(node["title"])][:1]
        fidelity = "先落回对象、证据、尺度、责任、观测、权力和行动上限，再决定判断档位。"
        misuse = "把术语当结论、把候选机制当事实、把开放断言当强判断。"
        lines.append(f"| {term} | {fidelity} | {misuse} | {linked_bundles(matched)} |")
    return "\n".join(lines)


def render_material_selection_map() -> str:
    lines = [
        "# CrossFrame v5 Material Selection Map",
        "",
        "本文件规定一次 CrossFrame 调用怎样从 v5 source modules 进入连读包、协议和输出模板。",
        "",
        "## 默认顺序",
        "",
        "1. 读 `references/v5-source-spine.md` 确认用户问题落在哪些 v5 原文节点。",
        "2. 读 `references/v5-section-digest-index.md` 获取相邻节点、摘要和不可误读边界。",
        "3. 读 `references/continuity-bundles.md`，先命中入口包，再展开“必须同读包”的闭包，最后才选择相邻候选包。",
        "4. 读取入口包和闭包包对应的 `references/continuity-bundles/v5/<bundle-id>.md`。",
        "5. 生成 `templates/read-state-capsule.md` 规定的 `v5-read-state-capsule`，把 source modules、入口包、必须同读闭包、相邻候选和降档边界交给下游。",
        "6. 进入 `references/read-routing-map.md` 指定的 protocol / worksheet / concept card / template。",
        "7. 输出前用 `worksheets/seven-gates-worksheet.md`、`worksheets/source-continuity-check.md` 与 `worksheets/source-anchor-integrity-check.md` 降档检查。",
        "",
        "## 闭包规则",
        "",
        "- 连读包不是按需孤立检索；命中一个入口包后，必须递归展开其“必须同读包”。",
        "- `3 个核心包 + 2 个辅助包` 只限制入口包和相邻候选包选择，不限制必须同读闭包。",
        "- 如果上下文不足以读取闭包中的硬依赖，不得跳过依赖包维持原判断，只能降档、暂停强判断或改为待补读。",
        "- 胶囊不是额外正文；它只记录本次读过什么、源锚点在哪里、哪些内容必须降档或标为本文推断。",
        "",
        "## 场景路由",
        "",
        "| 场景 | 必读材料 | 默认连读包 |",
        "| --- | --- | --- |",
    ]
    for route_name, materials, bundle_ids in ROUTES:
        lines.append(f"| {route_name} | {materials} | {linked_bundles(list(bundle_ids))} |")
    lines.extend(
        [
            "",
            "## 强制追加规则",
            "",
            "- 高责任文章、公共评论、组织处置、公开判断：追加 `v5-seven-gates-diagnosis-pack`、`v5-strong-judgment-eight-pack`、`v5-low-power-protection-pack`、`v5-evidence-downgrade-action-ceiling-pack`。",
            "- AI 报告、合规材料、机构自评：追加 `v5-ai-process-artifact-boundary-pack`，并声明过程性产物不能证明现实已经被验证。",
            "- 生命周期判断：追加 `v5-state-coordinate-lifecycle-pack`；必须写对象、层级、时间窗口、反向条件、暂停条件、撤回路径。",
            "- 亲密关系、无法退出、复杂创伤：追加 `v5-love-trapped-trauma-pack` 与 `v5-low-power-protection-pack`，先保护安全和退出空间。",
            "- 任何只读单张概念卡就要承担结论的情况：必须回到 `v5-core-concept-integrity-pack` 和相邻包。",
        ]
    )
    return "\n".join(lines)


def render_continuity_index() -> str:
    lines = [
        "# CrossFrame v5 Continuity Bundles",
        "",
        "本文件是 v5 连读包索引，不承载全部正文。每个包一个文件，路径为 `references/continuity-bundles/v5/<bundle-id>.md`。",
        "",
        "读取原则：先命中入口包，再递归展开必须同读包，形成连读闭包；相邻包只作为辅助候选。默认最多读取 3 个入口核心包和 2 个相邻辅助包，但这个上限不限制必须同读闭包。若上下文不足以读取闭包，降档，不把孤立概念卡或孤立包用于强判断。",
        "",
        "| 连读包 ID | 中文名 | 何时强制读取 | 必须同读闭包 | 相邻候选包 | 文件 |",
        "| --- | --- | --- | --- | --- | --- |",
    ]
    for bundle in V5_BUNDLES:
        closure = required_closure([bundle.id])
        lines.append(
            f"| `{bundle.id}` | {bundle.name} | {bundle.trigger} | {linked_bundles(closure)} | {linked_bundles(list(bundle.adjacent))} | `continuity-bundles/v5/{bundle.id}.md` |"
        )
    lines.extend(
        [
            "",
            "## 读取闭包执行",
            "",
            "1. 先从 `read-routing-map.md` 或 `v5-material-selection-map.md` 选出入口包。",
            "2. 对每个入口包读取本表“必须同读闭包”中的所有包；闭包可递归展开，直到没有新增硬依赖。",
            "3. 只有完成硬依赖闭包后，才从“相邻候选包”中按题材追加 0-2 个辅助包。",
            "4. 若硬依赖闭包太大或材料不足，不能只读入口包继续原判断；必须降档或暂停。",
            "",
            "## 输出影响总则",
            "",
            "- 连读包只约束结构、边界和判断档位；不得替代事实核验。",
            "- 高责任场景必须保留七闸、强判断八件套、低权力保护、证据降级和行动上限。",
            "- 文章输出继续由 `crossframe-essay` 生成结构洞察底稿和正文；review 只做隐藏质量闸摘要，不能吞掉正文。",
        ]
    )
    return "\n".join(lines)


def render_bundle_file(bundle: Bundle, nodes: list[dict[str, object]]) -> str:
    anchors = node_rows(nodes, bundle.id, limit=12)
    anchor_lines = [f"- `{node['id']}` {node['title']}（段落 {node['start']}-{node['end']}）" for node in anchors]
    if not anchor_lines:
        anchor_lines = ["- 未命中标题关键词时，按触发场景从 `v5-source-spine.md` 检索相邻节点。"]
    direct_required = list(REQUIRED_WITH.get(bundle.id, ()))
    closure = required_closure([bundle.id])
    return "\n".join(
        [
            f"# {bundle.name}",
            "",
            f"- ID：`{bundle.id}`",
            f"- 关键词：{'、'.join(bundle.keywords)}",
            "",
            "## 源锚点",
            "",
            *anchor_lines,
            "",
            "## 必须连读原因",
            "",
            bundle.reason,
            "",
            "## 触发场景",
            "",
            bundle.trigger,
            "",
            "## 必须同读材料",
            "",
            bundle.must_read,
            "",
            "## 必须同读包（硬约束）",
            "",
            "\n".join(f"- `{required}`" for required in direct_required) or "- 无",
            "",
            "## 必须同读闭包（递归展开）",
            "",
            "\n".join(f"- `{required}`" for required in closure) or "- 无",
            "",
            "## 硬失败",
            "",
            bundle.failure,
            "",
            "## 降档规则",
            "",
            bundle.downgrade,
            "",
            "## 相邻候选包（非硬约束）",
            "",
            "\n".join(f"- `{adjacent}`" for adjacent in bundle.adjacent),
            "",
            "## 输出影响",
            "",
            bundle.output,
            "",
            "## 输出自检",
            "",
            "- 是否先声明对象、事实、证据缺口和用途？",
            "- 是否明确本包只影响结构判断，不替代现实核验？",
            "- 是否写出降档、暂停或撤回条件？",
        ]
    )


def render_read_routing() -> str:
    lines = [
        "# 读取路由图",
        "",
        "本文件决定一次 CrossFrame 调用应该读取哪些材料。默认先读最少必要内容；当概念承担判断作用时，再加载完整概念卡、v5 source modules 和必要的连续联读包。",
        "",
        "当前权威源为 `v5.0`。默认连续性检查读取 `v5-source-spine.md`、`v5-section-digest-index.md`、`v5-coverage-map.md` 与 `v5-term-fidelity.md`。",
        "",
        "## v5 source modules",
        "",
        "- `v5-source-spine.md`：v5 原文标题顺序、段落范围、相邻关系、表格索引和默认连读包。",
        "- `v5-section-digest-index.md`：逐标题节点研究摘要、不可误读边界和相邻联读提醒。",
        "- `v5-coverage-map.md`：标题节点、连读包、协议、工作表和 skill 的覆盖关系。",
        "- `v5-term-fidelity.md`：v5 术语保真表。",
        "- `v5-material-selection-map.md`：从用户请求进入 v5 包、协议和模板的选择图。",
        "",
        "## v5 连读包索引",
        "",
        "| 连读包 | 中文名 | 触发场景 | 必须同读闭包 |",
        "| --- | --- | --- | --- |",
    ]
    for bundle in V5_BUNDLES:
        lines.append(f"| `{bundle.id}` | {bundle.name} | {bundle.trigger} | {linked_bundles(required_closure([bundle.id]))} |")
    lines.extend(
        [
            "",
        "## 基础路由",
        "",
        "| 用户请求 | 必读 | 连续联读包 |",
        "| --- | --- | --- |",
        ]
    )
    for route_name, materials, bundle_ids in ROUTES:
        lines.append(f"| {route_name} | {materials} | {linked_bundles(list(bundle_ids))} |")
    lines.extend(
        [
            "",
            "## 连续联读执行规则",
            "",
            "- 只要上表的连续联读包不是空，就先读 `references/continuity-bundles.md`，再读取对应 `references/continuity-bundles/v5/<bundle-id>.md`。",
            "- 连读包不是孤立按需检索。命中入口包后，必须递归展开该包的“必须同读闭包”；闭包中的包是硬依赖。",
            "- 默认最多读取 3 个入口核心包 + 2 个相邻辅助包；这个上限不限制必须同读闭包。若闭包无法读完，必须降档或暂停强判断。",
            "- 相邻候选包只在硬依赖闭包完成后按题材追加，不能替代必须同读包。",
            "- 深度、审计、高责任、公共制度、亲密关系、长期演化、框架治理和文章输出场景，必须用 `v5-source-spine.md` 或 `v5-section-digest-index.md` 检查相邻章节。",
            "- 若旧材料与 v5 的理解冲突，以 v5 为准。",
            "- 输出前使用 `worksheets/source-continuity-check.md`：若发现只读了单张概念卡，且本文件要求联读，必须补读或降档。",
            "- 输出前使用 `worksheets/source-anchor-integrity-check.md`：中心命题、机制候选、高风险概念和行动边界无法回指胶囊源锚点时，必须标为本文推断、表达转译或外部思想映射。",
            "- 下游 essay、专项 skill 和 review 默认复用 `v5-read-state-capsule`；只有胶囊缺少关键锚点、高责任审计或完整性检查失败时，才按具体锚点定向补读。",
            "- `templates/reasoning-outline-output.md` 中的本次连续联读包只列包名，不展开完整工作表。",
            "",
            "## 高风险概念触发",
            "",
            "| 触发词或判断动作 | 必读概念卡或协议 | 默认 v5 包 |",
            "| --- | --- | --- |",
            "| 证据、材料、报告、弱信号、AI 合规、自评 | `concept-cards/evidence-cost.md` | `v5-source-evidence-separation-pack`, `v5-ai-process-artifact-boundary-pack` |",
            "| 档位、能否强判断、能否处置、能否公开 | `concept-cards/judgment-grades.md` | `v5-evidence-downgrade-action-ceiling-pack`, `v5-strong-judgment-eight-pack` |",
            "| 尺度升维、换层解释、大局、历史、制度 | `concept-cards/scale-transfer.md` | `v5-cross-scale-context-translation-pack` |",
            "| 被诊断后变化、表演、反制、策略反应 | `concept-cards/reflexivity.md` | `v5-observation-reflexivity-release-pack` |",
            "| 爱、牺牲、忍耐、照护、开放行动 | `concept-cards/love-open-action.md` | `v5-love-trapped-trauma-pack` |",
            "| 公共承诺、平台治理、制度、分配回流 | `protocols/public-institution-protocol.md` | `v5-public-power-institution-pack` |",
            "| 阶段、生命周期、回退、混合阶段 | `protocols/lifecycle-diagnosis-protocol.md` | `v5-state-coordinate-lifecycle-pack` |",
            "| 工具化、公开发布、AI 工具、认证 | `references/framework-ontology-protection.md` | `v5-toolization-accessibility-release-pack` |",
            "",
            "## 输出路由",
            "",
            "- 默认输出使用 `templates/reasoning-outline-output.md` 作为前置提纲。",
            "- 快速诊断：`templates/quick-diagnosis-output.md`",
            "- 完整诊断：`templates/full-diagnosis-output.md`",
            "- 推演：`templates/inference-output.md`",
            "- 开放断言：`templates/open-assertion-output.md`",
            "- 概念解释：`templates/concept-explanation-output.md`",
            "- 强判断：`templates/strong-judgment-output.md`",
            "- 高反身性：`templates/high-reflexivity-output.md`",
            "- 亲密关系轻量入口：`templates/intimate-relationship-output.md`",
            "- 疗愈与转移：`templates/healing-transfer-output.md`",
            "- 公共制度专项：`templates/public-institution-output.md`",
            "- 框架边界：`templates/framework-boundary-output.md`",
            "- 生命周期：`templates/lifecycle-output.md`",
            "- 递进模式：`templates/progression-output.md`",
            "- 势场与自主解离：`templates/field-dissociation-output.md`",
            "- 治理连续性：`templates/governance-continuity-output.md`",
            "- 超大规模压力测试：`templates/large-scale-stress-output.md`",
            "- 对外表达翻译：`templates/expression-translation-output.md`",
            "",
            "输出前必须读或内化 `templates/user-facing-language.md` 的表达闸。",
        ]
    )
    return "\n".join(lines)


def render_read_state_capsule_template() -> str:
    return "\n".join(
        [
            "# v5-read-state-capsule",
            "",
            "本模板用于 `crossframe` 核心层在完成 v5 source modules、入口连读包、必须同读闭包和相邻候选选择后，生成一次性的源结构摘要。它不是文章正文，也不是完整原文摘录；它是传给专项 skill、`crossframe-essay` 和 `crossframe-review` 的读态记录。",
            "",
            "## 使用位置",
            "",
            "- `crossframe-suite` 只传入 `selection_state` 和 `workflow_state`，不生成胶囊。",
            "- `crossframe` 在读取 `v5-source-spine.md`、`v5-section-digest-index.md`、`v5-material-selection-map.md`、`continuity-bundles.md` 和必要包文件后生成胶囊。",
            "- `crossframe-essay`、公共/组织/辩论/读书等专项 skill、`crossframe-review` 默认复用胶囊；不得各自重新发明源路由。",
            "- 只有胶囊缺少关键锚点、高责任审计需要复核、或 `source-anchor-integrity-check.md` 失败时，才按具体 source module、章节锚点或连读包定向补读。",
            "",
            "## 胶囊字段",
            "",
            "```text",
            "v5-read-state-capsule",
            "- selection_state：",
            "  - output_mode：",
            "  - role：",
            "  - topic_sensitivity：",
            "  - voice_mode：",
            "  - user_closed_article_layer：",
            "  - user_closed_review：",
            "- workflow_state：",
            "- user_task：",
            "- v5_source_modules：",
            "  - source_module_id / V5-H 锚点：",
            "  - 源范围：",
            "  - 触发理由：",
            "  - 必读相邻模块：",
            "  - 本次使用的模块摘要：",
            "  - 降档边界：",
            "- v5_continuity_bundles：",
            "  - 入口包 ID：",
            "  - 触发理由：",
            "  - 源锚点 / 章节范围：",
            "  - 本次使用的源结构摘要：",
            "  - 不可单读风险：",
            "  - 降档规则：",
            "- required_closure：",
            "  - 必须同读闭包：",
            "  - 闭包是否读完：",
            "  - 未读完时的降档决定：",
            "- adjacent_candidates：",
            "  - 已追加的相邻候选包：",
            "  - 未追加但需注意的相邻约束：",
            "- source_grounding：",
            "  - 中心命题可用源锚点：",
            "  - 机制候选可用源锚点：",
            "  - 高风险概念可用源锚点：",
            "  - 行动 / 边界可用源锚点：",
            "  - 文章类型转译可用源锚点：",
            "  - 写作技法不能越过的源边界：",
            "  - 需要标为“本文推断 / 表达转译 / 外部思想映射”的内容：",
            "- downstream_read_policy：",
            "  - 默认只读本胶囊和本 skill 协议：",
            "  - 禁止重复整块读取的材料：",
            "  - 允许定向补读的条件：",
            "- integrity_risks：",
            "```",
            "",
            "## 写法要求",
            "",
            "- 胶囊必须先列 `v5_source_modules`，再列 `v5_continuity_bundles`、`required_closure` 和 `adjacent_candidates`。",
            "- 每个 source module、连读包和闭包包只写本次相关摘要，不复制大段原文、完整索引或完整包说明。",
            "- 每条承担判断作用的中心命题、机制候选、高风险概念、行动边界、文章类型转译和写作技法使用，至少关联一个源锚点、章节范围或连读包。",
            "- 不能从源锚点推出的内容，必须标成“本文推断”“表达转译”或“外部思想映射”。",
            "- 胶囊应尽量短；普通文章任务控制在 800-1500 中文字，高责任审计可更长。",
            "- 胶囊可在结构洞察底稿中摘要显示；正文不得出现 `v5-read-state-capsule` 标题或流程说明。",
            "",
            "## 定向补读条件",
            "",
            "- 胶囊缺少本次中心命题、机制候选或行动边界需要的源锚点。",
            "- 高责任、公开判断、现实处置、AI/过程性产物、无法退出主体、生命周期判断或用户要求审计源结构。",
            "- `source-anchor-integrity-check.md` 发现漏读前置包、只读单卡、闭包不完整或源锚点不足。",
            "- `crossframe-review` 需要定位连续性保真失败。",
            "",
            "补读时只打开相关锚点、章节摘要或包文件，不整块吞入全量索引。",
        ]
    )


def render_source_anchor_integrity_check() -> str:
    return "\n".join(
        [
            "# 源锚点完整性检查表",
            "",
            "本工作表用于输出前检查 `v5-read-state-capsule` 是否足以支撑最终判断。它和 `source-continuity-check.md` 配套：连续性检查负责是否读对包，源锚点检查负责最终判断能否回指本次已读源结构。",
            "",
            "默认不完整展示给普通用户；在结构洞察底稿中只摘要“源锚点覆盖 / 无法回指内容 / 降档决定”。",
            "",
            "## 1. 胶囊存在性",
            "",
            "| 项目 | 填写 |",
            "| --- | --- |",
            "| 是否存在 `v5-read-state-capsule` | 是 / 否 |",
            "| 胶囊是否先列 `v5_source_modules` | 是 / 否 |",
            "| 胶囊是否列出 `v5_continuity_bundles` | 是 / 否 |",
            "| 胶囊是否列出 `required_closure` | 是 / 否 |",
            "| 胶囊是否列出 `adjacent_candidates` | 是 / 否 |",
            "| 胶囊是否写明降档边界 | 是 / 否 |",
            "",
            "若胶囊缺失，回到 `crossframe` 核心层补生成；不得让 `crossframe-essay` 或 `crossframe-review` 各自重新整块读取源索引。",
            "",
            "## 2. 源锚点覆盖",
            "",
            "| 输出部位 | 必须回指 | 无法回指时处理 |",
            "| --- | --- | --- |",
            "| 中心命题 | 胶囊中的 V5-H 锚点、章节范围或入口连读包 | 删除、降档，或标为本文推断 |",
            "| 机制候选 | 胶囊中的源锚点、必须同读闭包或相邻约束 | 降为候选，不写成框架原义 |",
            "| 高风险概念 | 对应概念卡 + 连读包 + 必须同读闭包 | 补读或改为表达转译 |",
            "| 行动边界 | 七闸、证据降级、行动上限、低权力保护相关锚点 | 降为低条件行动或暂停建议 |",
            "| 文章类型转译 | 底稿中的事实边界、判断档位和胶囊源边界 | 只改表达，不改判断 |",
            "| 写作技法 | 胶囊允许的表达边界和底稿中心命题 | 技法句式不得新增事实或强判断 |",
            "| 概念上升 / 经典互文 | 源锚点 + 明确“思想映射”边界 | 标为外部思想映射，不冒充框架原义 |",
            "",
            "## 3. 必须执行场景",
            "",
            "以下场景必须执行源锚点完整性检查：",
            "",
            "- 高责任、公开判断、资格、名誉、权利、资源、处分、公共记忆。",
            "- 公共议题、平台治理、政策、机构责任、公共评论。",
            "- AI 报告、合规材料、机构自评、流程表、模型诊断或过程性产物。",
            "- 生命周期、趋势推演、阶段判断、长期演化。",
            "- 亲密关系、无法退出主体、复杂创伤、低权力主体。",
            "- 框架治理、证伪、工具化、公开发布、商业化。",
            "- 任何文章输出，尤其使用文章类型选择器和写作技法时。",
            "",
            "## 4. 失败处理",
            "",
            "| 失败类型 | 处理 |",
            "| --- | --- |",
            "| 无胶囊 | 回到 `crossframe` 生成胶囊，不继续成文或评审 |",
            "| 无 source modules | 回到 `v5-material-selection-map.md` 定向补读 |",
            "| 无入口包或闭包不完整 | 补读闭包；不能补读则降档 |",
            "| 中心命题无源锚点 | 删除、降档，或标为本文推断 |",
            "| 机制候选无源锚点 | 保留为候选，不得写成框架结论 |",
            "| 行动建议无锚点 | 降为低风险观察、补证清单或暂停建议 |",
            "| 写作技法新增判断 | 删除新增判断，回到底稿和胶囊边界 |",
            "| 经典互文接管命题 | 标为外部思想映射，不能压过事实边界 |",
            "",
            "## 5. 底稿显示",
            "",
            "结构洞察底稿只需显示：",
            "",
            "```text",
            "- 读态胶囊摘要：...",
            "- 源锚点覆盖：中心命题 / 机制候选 / 高风险概念 / 行动边界",
            "- 无法回指内容：本文推断 / 表达转译 / 外部思想映射",
            "- 降档决定：...",
            "```",
            "",
            "正文不得显示本工作表标题；只把检查结果转化为边界段、撤回条件和来源透明说明。",
        ]
    )


def generate(repo: Path, source_docx: Path) -> None:
    crossframe = repo / "skills" / "crossframe"
    refs = crossframe / "references"
    templates = crossframe / "templates"
    worksheets = crossframe / "worksheets"
    bundles_dir = refs / "continuity-bundles" / "v5"
    nodes, tables = extract_nodes(source_docx)
    write(refs / "v5-source-spine.md", render_source_spine(nodes, tables, source_docx))
    write(refs / "v5-section-digest-index.md", render_digest(nodes, source_docx))
    write(refs / "v5-coverage-map.md", render_coverage(nodes, tables))
    write(refs / "v5-term-fidelity.md", render_term_fidelity(nodes))
    write(refs / "v5-material-selection-map.md", render_material_selection_map())
    write(refs / "continuity-bundles.md", render_continuity_index())
    write(refs / "read-routing-map.md", render_read_routing())
    write(templates / "read-state-capsule.md", render_read_state_capsule_template())
    write(worksheets / "source-anchor-integrity-check.md", render_source_anchor_integrity_check())
    for bundle in V5_BUNDLES:
        write(bundles_dir / f"{bundle.id}.md", render_bundle_file(bundle, nodes))
    print(f"generated v5 source continuity: nodes={len(nodes)} tables={len(tables)} bundles={len(V5_BUNDLES)}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate CrossFrame source continuity files from a DOCX.")
    parser.add_argument("--version", default="v5")
    parser.add_argument("--source-docx", default=V5_DOCX)
    parser.add_argument("--repo", default=".")
    args = parser.parse_args()

    version = args.version.lower()
    if version != "v5":
        raise SystemExit("this generator is versioned for v5; pass --version v5")
    repo = Path(args.repo).resolve()
    source_docx = Path(args.source_docx)
    if not source_docx.exists():
        raise SystemExit(f"missing source docx: {source_docx}")
    generate(repo, source_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
